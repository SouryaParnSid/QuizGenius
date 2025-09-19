"""
Document Processing Service

Service for processing various document types, chunking text,
and preparing documents for vector storage.
"""

import logging
import re
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import hashlib
from datetime import datetime

# Document processing imports
import pypdf
from docx import Document as DocxDocument

# LangChain text splitters
from langchain.text_splitter import (
    RecursiveCharacterTextSplitter,
    TokenTextSplitter,
    MarkdownTextSplitter
)

from .config import RAGConfig
from .vector_store import Document

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing and chunking documents."""
    
    def __init__(self, config: RAGConfig):
        """Initialize the document processor."""
        self.config = config
        
        # Initialize text splitters
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.chunk_size,
            chunk_overlap=config.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        self.token_splitter = TokenTextSplitter(
            chunk_size=config.max_tokens_per_chunk,
            chunk_overlap=config.chunk_overlap // 4  # Smaller overlap for tokens
        )
        
        self.markdown_splitter = MarkdownTextSplitter(
            chunk_size=config.chunk_size,
            chunk_overlap=config.chunk_overlap
        )
    
    def extract_text_from_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF file."""
        try:
            text = ""
            metadata = {
                "file_type": "pdf",
                "file_path": file_path,
                "pages": 0,
                "extraction_method": "pypdf"
            }
            
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                metadata["pages"] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text += f"\n\n--- Page {page_num + 1} ---\n\n"
                            text += page_text
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
            
            metadata["text_length"] = len(text)
            metadata["success"] = len(text) > 0
            
            return text.strip(), metadata
            
        except Exception as e:
            logger.error(f"Failed to extract text from PDF {file_path}: {e}")
            return "", {"error": str(e), "file_type": "pdf", "success": False}
    
    def extract_text_from_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n\n".join(text_parts)
            
            metadata = {
                "file_type": "docx",
                "file_path": file_path,
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "text_length": len(text),
                "success": len(text) > 0
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            return "", {"error": str(e), "file_type": "docx", "success": False}
    
    def extract_text_from_txt(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            metadata = {
                "file_type": "txt",
                "file_path": file_path,
                "text_length": len(text),
                "lines": len(text.split('\n')),
                "success": len(text) > 0
            }
            
            return text, metadata
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                
                metadata = {
                    "file_type": "txt",
                    "file_path": file_path,
                    "text_length": len(text),
                    "lines": len(text.split('\n')),
                    "encoding": "latin-1",
                    "success": len(text) > 0
                }
                
                return text, metadata
                
            except Exception as e:
                logger.error(f"Failed to extract text from TXT {file_path}: {e}")
                return "", {"error": str(e), "file_type": "txt", "success": False}
        
        except Exception as e:
            logger.error(f"Failed to extract text from TXT {file_path}: {e}")
            return "", {"error": str(e), "file_type": "txt", "success": False}
    
    def extract_text_from_markdown(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from Markdown file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            metadata = {
                "file_type": "markdown",
                "file_path": file_path,
                "text_length": len(text),
                "lines": len(text.split('\n')),
                "success": len(text) > 0
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"Failed to extract text from Markdown {file_path}: {e}")
            return "", {"error": str(e), "file_type": "markdown", "success": False}
    
    def extract_text_from_file(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from file based on extension."""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return self.extract_text_from_pdf(str(file_path))
        elif extension == '.docx':
            return self.extract_text_from_docx(str(file_path))
        elif extension == '.txt':
            return self.extract_text_from_txt(str(file_path))
        elif extension in ['.md', '.markdown']:
            return self.extract_text_from_markdown(str(file_path))
        else:
            error_msg = f"Unsupported file type: {extension}"
            logger.error(error_msg)
            return "", {"error": error_msg, "file_type": extension, "success": False}
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive newlines
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Remove special characters that might interfere with processing
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f]', '', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def chunk_text(
        self,
        text: str,
        chunking_strategy: str = "recursive",
        custom_separators: Optional[List[str]] = None
    ) -> List[str]:
        """Chunk text into smaller pieces."""
        if not text or not text.strip():
            return []
        
        # Clean text first
        text = self.clean_text(text)
        
        try:
            if chunking_strategy == "recursive":
                if custom_separators:
                    splitter = RecursiveCharacterTextSplitter(
                        chunk_size=self.config.chunk_size,
                        chunk_overlap=self.config.chunk_overlap,
                        separators=custom_separators
                    )
                    chunks = splitter.split_text(text)
                else:
                    chunks = self.text_splitter.split_text(text)
            
            elif chunking_strategy == "token":
                chunks = self.token_splitter.split_text(text)
            
            elif chunking_strategy == "markdown":
                chunks = self.markdown_splitter.split_text(text)
            
            elif chunking_strategy == "paragraph":
                # Split by paragraphs
                paragraphs = text.split('\n\n')
                chunks = []
                current_chunk = ""
                
                for paragraph in paragraphs:
                    if len(current_chunk) + len(paragraph) <= self.config.chunk_size:
                        current_chunk += paragraph + "\n\n"
                    else:
                        if current_chunk:
                            chunks.append(current_chunk.strip())
                        current_chunk = paragraph + "\n\n"
                
                if current_chunk:
                    chunks.append(current_chunk.strip())
            
            else:
                logger.warning(f"Unknown chunking strategy: {chunking_strategy}, using recursive")
                chunks = self.text_splitter.split_text(text)
            
            # Filter out very small chunks
            min_chunk_size = 50  # Minimum 50 characters
            chunks = [chunk for chunk in chunks if len(chunk) >= min_chunk_size]
            
            return chunks
            
        except Exception as e:
            logger.error(f"Failed to chunk text: {e}")
            return [text]  # Return original text as single chunk
    
    def process_file(
        self,
        file_path: str,
        chunking_strategy: str = "recursive",
        custom_metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """Process a file and return a list of Document objects."""
        file_path = Path(file_path)
        
        # Validate file
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return []
        
        if file_path.suffix.lower() not in self.config.supported_file_types:
            logger.error(f"Unsupported file type: {file_path.suffix}")
            return []
        
        if file_path.stat().st_size > self.config.max_file_size:
            logger.error(f"File too large: {file_path}")
            return []
        
        # Extract text
        text, extraction_metadata = self.extract_text_from_file(str(file_path))
        
        if not text or not extraction_metadata.get("success", False):
            logger.error(f"Failed to extract text from {file_path}")
            return []
        
        # Chunk text
        chunks = self.chunk_text(text, chunking_strategy)
        
        if not chunks:
            logger.warning(f"No chunks generated from {file_path}")
            return []
        
        # Create documents
        documents = []
        file_hash = self._get_file_hash(str(file_path))
        
        for i, chunk in enumerate(chunks):
            metadata = {
                "source_file": str(file_path),
                "file_name": file_path.name,
                "file_type": file_path.suffix.lower(),
                "file_hash": file_hash,
                "chunk_index": i,
                "total_chunks": len(chunks),
                "chunk_size": len(chunk),
                "chunking_strategy": chunking_strategy,
                "processed_at": datetime.now().isoformat(),
                **extraction_metadata
            }
            
            # Add custom metadata if provided
            if custom_metadata:
                metadata.update(custom_metadata)
            
            doc_id = f"{file_hash}_chunk_{i}"
            document = Document(content=chunk, metadata=metadata, doc_id=doc_id)
            documents.append(document)
        
        logger.info(f"Processed {file_path}: {len(chunks)} chunks created")
        return documents
    
    def process_text(
        self,
        text: str,
        source_name: str = "text_input",
        chunking_strategy: str = "recursive",
        custom_metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """Process raw text and return a list of Document objects."""
        if not text or not text.strip():
            return []
        
        # Chunk text
        chunks = self.chunk_text(text, chunking_strategy)
        
        if not chunks:
            return []
        
        # Create documents
        documents = []
        text_hash = hashlib.md5(text.encode()).hexdigest()
        
        for i, chunk in enumerate(chunks):
            metadata = {
                "source": source_name,
                "source_type": "text",
                "text_hash": text_hash,
                "chunk_index": i,
                "total_chunks": len(chunks),
                "chunk_size": len(chunk),
                "chunking_strategy": chunking_strategy,
                "processed_at": datetime.now().isoformat(),
                "original_text_length": len(text)
            }
            
            # Add custom metadata if provided
            if custom_metadata:
                metadata.update(custom_metadata)
            
            doc_id = f"{text_hash}_chunk_{i}"
            document = Document(content=chunk, metadata=metadata, doc_id=doc_id)
            documents.append(document)
        
        logger.info(f"Processed text input: {len(chunks)} chunks created")
        return documents
    
    def _get_file_hash(self, file_path: str) -> str:
        """Generate hash for file content."""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    
    def get_supported_types(self) -> List[str]:
        """Get list of supported file types."""
        return list(self.config.supported_file_types)
    
    def validate_file(self, file_path: str) -> Tuple[bool, str]:
        """Validate if file can be processed."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            return False, "File not found"
        
        if file_path.suffix.lower() not in self.config.supported_file_types:
            return False, f"Unsupported file type: {file_path.suffix}"
        
        if file_path.stat().st_size > self.config.max_file_size:
            return False, f"File too large (max: {self.config.max_file_size / 1024 / 1024:.1f}MB)"
        
        return True, "File is valid"
    
    def get_text_stats(self, text: str) -> Dict[str, Any]:
        """Get statistics about text content."""
        lines = text.split('\n')
        words = text.split()
        sentences = re.split(r'[.!?]+', text)
        
        return {
            "character_count": len(text),
            "word_count": len(words),
            "line_count": len(lines),
            "sentence_count": len([s for s in sentences if s.strip()]),
            "paragraph_count": len([p for p in text.split('\n\n') if p.strip()]),
            "avg_words_per_sentence": len(words) / max(len(sentences), 1),
            "avg_chars_per_word": len(text) / max(len(words), 1)
        }
